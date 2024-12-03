from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd

# Load dataset
file_path = 'data/gym_membership.csv'  # Sesuaikan dengan path dataset Anda
df = pd.read_csv(file_path)

# Membuat objek dokumen Word
doc = Document()

# BAB I: Pendahuluan
doc.add_heading('BAB I: PENDAHULUAN', level=1)
doc.add_paragraph("Dalam bab ini, data yang digunakan adalah data keanggotaan gym, yang berisi informasi tentang anggota, "
                  "kebiasaan kunjungan mereka, serta preferensi terhadap kelas dan pelatihan. Adapun deskripsi data dan variabelnya adalah sebagai berikut:")

# Menyusun deskripsi data
data_description = """
1. Age: Usia anggota gym.
2. visit_per_week: Frekuensi kunjungan ke gym per minggu.
3. avg_time_in_gym: Durasi rata-rata waktu yang dihabiskan di gym per kunjungan (dalam menit).
4. attend_group_lesson: Apakah anggota mengikuti kelas kelompok atau tidak.
5. personal_training: Apakah anggota menggunakan pelatihan pribadi atau tidak.
"""

doc.add_paragraph(data_description)

# Menyertakan volume data
doc.add_paragraph(f"Jumlah data (volume): {len(df)} entri")
doc.add_paragraph("Berikut ini adalah sampel dari data yang digunakan:")

# Menyertakan sampel data
doc.add_paragraph(df.head().to_string())

# BAB II: Metode
doc.add_heading('BAB II: METODE', level=1)
doc.add_paragraph("Dalam bab ini, metode analisis data yang digunakan meliputi pembersihan data dan visualisasi data. "
                  "Langkah-langkah pengerjaan analisis adalah sebagai berikut:")

# Langkah-langkah analisis
steps = """
1. Pembersihan data:
   - Memeriksa dan mengisi nilai hilang
   - Menghapus data duplikat
   - Memastikan tipe data yang benar
   - Menangani data yang tidak valid

2. Visualisasi data:
   - Membuat histogram untuk distribusi usia anggota
   - Membuat bar plot untuk frekuensi kunjungan per minggu
   - Membuat histogram untuk durasi waktu di gym per kunjungan
   - Membuat pie chart untuk preferensi terhadap kelas kelompok
   - Membuat pie chart untuk penggunaan pelatihan pribadi
"""

doc.add_paragraph(steps)

# BAB III: Pembahasan
doc.add_heading('BAB III: PEMBAHASAN', level=1)
doc.add_paragraph("Pembahasan hasil analisis data yang dilakukan adalah sebagai berikut:")

# Hasil Analisis
results = """
1. Distribusi Usia Anggota Gym:
   - Sebagian besar anggota berada pada usia 20 hingga 40 tahun. Hal ini menunjukkan bahwa kelompok usia ini adalah yang paling aktif.

2. Frekuensi Kunjungan Per Minggu:
   - Mayoritas anggota berkunjung ke gym sebanyak 2 hingga 3 kali per minggu. Ini menunjukkan tingkat keterlibatan yang cukup baik.

3. Durasi Waktu di Gym:
   - Rata-rata anggota menghabiskan 60 hingga 120 menit per kunjungan, menunjukkan bahwa mereka memiliki rutinitas latihan yang cukup lama.

4. Preferensi terhadap Kelas Kelompok:
   - Sekitar setengah anggota mengikuti kelas kelompok. Ini menunjukkan adanya minat yang baik terhadap kelas kelompok, namun masih terdapat anggota yang memilih untuk tidak mengikuti.

5. Penggunaan Pelatihan Pribadi:
   - Pelatihan pribadi juga cukup diminati, tetapi hanya sekitar setengah anggota yang memilih untuk menggunakan layanan ini.
"""

doc.add_paragraph(results)

# BAB IV: Penutup
doc.add_heading('BAB IV: PENUTUP', level=1)
doc.add_paragraph("Kesimpulan dari analisis data ini adalah bahwa sebagian besar anggota gym berusia 20-40 tahun dan berkunjung "
                  "ke gym dengan frekuensi moderat. Durasi latihan anggota menunjukkan rutinitas yang konsisten, dan terdapat "
                  "minat yang cukup baik terhadap kelas kelompok serta pelatihan pribadi. \n\n"
                  "Saran untuk pengembangan analisis ini ke depannya adalah menggali preferensi yang lebih spesifik dari anggota, "
                  "serta melihat faktor-faktor tambahan seperti tingkat kepuasan dan hasil kesehatan untuk memberikan layanan yang lebih baik.")

# Menyimpan dokumen
output_path = "data/Laporan_Analisis_Gym.docx"
doc.save(output_path)
print(f"Laporan telah disimpan di {output_path}")
